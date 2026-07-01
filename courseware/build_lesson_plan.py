#!/usr/bin/env python3
"""
3-day Lesson Plan (DOCX) for
"Microsoft Copilot Studio & Power Automate for Business Workflow Automation".
CLIENT TRAINING — de-branded (no Tertiary logo/name/UEN/copyright).
Daily window 9:00am - 5:00pm (8 hours incl. 1-hour lunch + two tea breaks).
Writes: courseware/Lesson Plan - <course>.docx
"""
import os, sys
SKILL = "/Users/alfredang/.claude/skills/tertiary-lesson-plan"
sys.path.insert(0, SKILL)
import prodoc
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TITLE = "Microsoft Copilot Studio & Power Automate for Business Workflow Automation"
VERSION = "1.0"
COURSE_CODE = "TGS-2022017524"
VERSIONS = [["1.0", "24 Jun 2026", "Initial release — 3-day lesson plan (9:00am-5:00pm).",
             "Course Development Team"]]

BRAND = RGBColor(0x1F,0x6F,0xEB); DARK = RGBColor(0x16,0x1B,0x26); GREY = RGBColor(0x55,0x5B,0x66)
CENTER = WD_ALIGN_PARAGRAPH.CENTER

def _cline(doc, text, size, bold=False, color=DARK, before=0, after=4):
    p = doc.add_paragraph(); p.alignment = CENTER
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = color; r.font.name = "Arial"
    return p
def add_cover(doc):
    for _ in range(4): doc.add_paragraph()
    _cline(doc, "LESSON PLAN", 26, bold=True, color=BRAND, after=12)
    _cline(doc, "For", 12, color=GREY, after=8)
    _cline(doc, TITLE, 20, bold=True, color=DARK, after=14)
    _cline(doc, f"Course Code: {COURSE_CODE}", 12, color=GREY, after=6)
    _cline(doc, "Duration: 3 Days  ·  9:00am – 5:00pm", 12, color=GREY, after=18)
    _cline(doc, f"Version {VERSION}", 12, bold=True, color=BRAND)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
def add_footer(doc):
    sec = doc.sections[0]; f = sec.footer; f.is_linked_to_previous = False
    p = f.paragraphs[0]; p.alignment = CENTER; p.text = ""
    r = p.add_run("Page "); r.font.size = Pt(9); r.font.color.rgb = GREY
    prodoc._field(p, "PAGE", "1").font.size = Pt(9)
    r2 = p.add_run(" of "); r2.font.size = Pt(9); r2.font.color.rgb = GREY
    prodoc._field(p, "NUMPAGES", "1").font.size = Pt(9)
    sp = f.add_paragraph(); sp.alignment = CENTER
    sr = sp.add_run(TITLE + " — Lesson Plan"); sr.font.size = Pt(7.5); sr.font.color.rgb = GREY

def _shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hexc); tcPr.append(shd)

def _borders(t, color="7A8190", sz=6):
    """Apply explicit single borders (outer + inner grid) so every cell edge is visible."""
    tblPr = t._tbl.tblPr
    for old in tblPr.findall(qn("w:tblBorders")): tblPr.remove(old)
    b = OxmlElement("w:tblBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"),"single"); e.set(qn("w:sz"),str(sz))
        e.set(qn("w:space"),"0"); e.set(qn("w:color"),color)
        b.append(e)
    tblPr.append(b)

def heading(doc, text, lvl=1):
    doc.add_paragraph(style=f"Heading {lvl}").add_run(text)
def para(doc, text):
    p = doc.add_paragraph(); p.add_run(text); return p
def bullets(doc, items):
    for it in items: doc.add_paragraph(style="List Bullet").add_run(it)

def info_table(doc, rows):
    t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in rows:
        c = t.add_row().cells
        c[0].text = ""; rr = c[0].paragraphs[0].add_run(k); rr.bold = True; rr.font.size = Pt(10); _shade(c[0], "EAF1FF")
        c[1].text = ""; c[1].paragraphs[0].add_run(v).font.size = Pt(10)
    _borders(t)

def schedule_table(doc, rows):
    """rows: (time, activity, duration, kind, slides) ; kind in {'topic','lab','break',''}"""
    t = doc.add_table(rows=0, cols=4); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.add_row().cells
    for i, h in enumerate(["Time", "Topic / Activity", "Duration", "Slides"]):
        hdr[i].text = ""; rr = hdr[i].paragraphs[0].add_run(h)
        rr.bold = True; rr.font.size = Pt(9.5); rr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); _shade(hdr[i], "1F6FEB")
    for time, act, dur, kind, slides in rows:
        cells = t.add_row().cells
        vals = [time, act, dur, slides]
        tint = {"break":"FFF4E5", "topic":"EAF1FF", "lab":"E8F7EE"}.get(kind)
        for i, v in enumerate(vals):
            cells[i].text = ""; pp = cells[i].paragraphs[0]
            rr = pp.add_run(v); rr.font.size = Pt(9.5)
            if i == 1 and kind in ("topic","lab"): rr.bold = True
            if tint: _shade(cells[i], tint)
    # widths
    for row in t.rows:
        row.cells[0].width = Pt(82); row.cells[2].width = Pt(54); row.cells[3].width = Pt(58)
    _borders(t)

# ---------------------------------------------------------------- schedules
DAY1 = [
    ("9:00 – 9:30",  "Welcome, introductions, environment & accounts check (Lab 0)", "30 min", "topic", "1–5"),
    ("9:30 – 10:15", "Module 1: Introduction to Workflow Automation — triggers, actions, outputs, steps", "45 min", "topic", "6–17"),
    ("10:15 – 10:45","Module 2: Introduction to Power Automate — flow types, common triggers, connectors", "30 min", "topic", "18–30"),
    ("10:45 – 11:00","Tea break", "15 min", "break", "—"),
    ("11:00 – 12:00","Lab 1: Automated Email Workflow (manual trigger → Send an email)", "60 min", "lab", "31–32"),
    ("12:00 – 12:30","Lab 2: Excel Data Logging Workflow — build the table & flow", "30 min", "lab", "33"),
    ("12:30 – 1:30", "Lunch", "60 min", "break", "—"),
    ("1:30 – 2:15",  "Lab 2 (cont.): map columns, fx date expression, test & verify rows", "45 min", "lab", "33"),
    ("2:15 – 3:15",  "Lab 3: Simple Approval Workflow — approvals & conditions", "60 min", "lab", "34"),
    ("3:15 – 3:30",  "Tea break", "15 min", "break", "—"),
    ("3:30 – 4:15",  "Lab 4: Scheduled Trigger Workflow — Recurrence & reminders", "45 min", "lab", "35"),
    ("4:15 – 5:00",  "Lab 5: Form Submission Workflow + Day 1 recap & Q&A", "45 min", "lab", "36–37"),
]
DAY2 = [
    ("9:00 – 9:15",  "Day 1 recap & Q&A", "15 min", "topic", "37–39"),
    ("9:15 – 10:00", "Module 3: Building Business Agents with Copilot Studio — prompt design for structured output", "45 min", "topic", "40–49"),
    ("10:00 – 10:45","Lab 6: Create Your First Agent — instructions, knowledge, Test pane", "45 min", "lab", "55–56"),
    ("10:45 – 11:00","Tea break", "15 min", "break", "—"),
    ("11:00 – 11:45","Lab 7: Add Knowledge to Your Agent (Grounding / RAG)", "45 min", "lab", "57"),
    ("11:45 – 12:30","Lab 8: Add Tools and Actions — connectors & flows as tools", "45 min", "lab", "58"),
    ("12:30 – 1:30", "Lunch", "60 min", "break", "—"),
    ("1:30 – 2:30",  "Lab 9: Sales Enquiry Assistant — capture structured data", "60 min", "lab", "59"),
    ("2:30 – 3:15",  "Lab 10: Procurement Request Workflow — agent calls a flow", "45 min", "lab", "50–54, 60"),
    ("3:15 – 3:30",  "Tea break", "15 min", "break", "—"),
    ("3:30 – 4:15",  "Lab 10 (cont.): log + notify, wire inputs, test end-to-end", "45 min", "lab", "60"),
    ("4:15 – 5:00",  "Lab 11: Automated Response Generation (AI prompts) + recap", "45 min", "lab", "61–63"),
]
DAY3 = [
    ("9:00 – 9:15",  "Day 2 recap & Q&A", "15 min", "topic", "63–64"),
    ("9:15 – 9:45",  "Module 4: End-to-End Orchestration Concepts", "30 min", "topic", "65–71"),
    ("9:45 – 10:45", "Lab 12: Email Enquiry → Excel Logging → Notification", "60 min", "lab", "72"),
    ("10:45 – 11:00","Tea break", "15 min", "break", "—"),
    ("11:00 – 12:00","Lab 13: Invoice Upload → Approval Workflow (file trigger)", "60 min", "lab", "73"),
    ("12:00 – 12:30","Lab 14: Purchase Request → Manager Approval — build", "30 min", "lab", "74"),
    ("12:30 – 1:30", "Lunch", "60 min", "break", "—"),
    ("1:30 – 2:15",  "Lab 14 (cont.): threshold condition, notify, test all paths", "45 min", "lab", "74"),
    ("2:15 – 3:00",  "Lab 15: Order Processing (Agent + Flow) — confirm, log, restock alert", "45 min", "lab", "75–76"),
    ("3:00 – 3:15",  "Tea break", "15 min", "break", "—"),
    ("3:15 – 3:30",  "Module 5: Business Workflow Workshop briefing", "15 min", "topic", "77–83"),
    ("3:30 – 4:45",  "Lab 16: Capstone Workshop — design, build & present your own workflow", "75 min", "lab", "84–85"),
    ("4:45 – 5:00",  "Course wrap-up, feedback & where to go next", "15 min", "topic", "86–94"),
]

def total(rows):
    return sum(int(r[2].split()[0]) for r in rows)
for nm, rows in [("Day 1",DAY1),("Day 2",DAY2),("Day 3",DAY3)]:
    assert total(rows) == 480, f"{nm} = {total(rows)} min (expected 480)"

# ---------------------------------------------------------------- build doc
doc = Document()
nrm = doc.styles["Normal"]; nrm.font.name = "Arial"; nrm.font.size = Pt(11)
prodoc.style_headings(doc)
add_cover(doc)
prodoc.add_version_control(doc, VERSIONS)
prodoc.add_toc(doc, levels="1-1")

heading(doc, "Course Overview", 1)
para(doc, "This 3-day, hands-on course teaches participants to automate real business processes by combining "
          "Microsoft Power Automate flows with AI agents built in Microsoft Copilot Studio. Every topic follows "
          "a concept → demonstration → hands-on lab pattern; participants finish with working automations and a "
          "capstone project of their own.")
info_table(doc, [
    ("Duration", "3 Days (9:00am – 5:00pm), 8 hours per day incl. 1-hour lunch"),
    ("Delivery", "Instructor-led, hands-on (physical / virtual)"),
    ("Audience", "Business and operations staff who want to automate repetitive work; no coding required"),
    ("Prerequisites", "Basic Microsoft 365 familiarity (Outlook, Excel); a Power Platform environment (see Lab 0)"),
    ("Labs", "17 step-by-step labs across 3 days, plus 5 concept modules"),
])

heading(doc, "Learning Outcomes", 1)
para(doc, "By the end of the course, participants will be able to:")
bullets(doc, [
    "Explain business workflow automation and the Trigger → Actions → Output model.",
    "Build Power Automate flows that send emails, log data to Excel, run approvals, and use scheduled and form triggers.",
    "Create business agents in Copilot Studio, ground them with knowledge (RAG), and give them tools.",
    "Connect agents to Power Automate flows and pass structured data between them.",
    "Orchestrate complete end-to-end workflows for sales, finance, procurement, and order processing.",
    "Design, build, and present an original end-to-end automation in a capstone project.",
])

heading(doc, "Daily Schedule", 1)
para(doc, "The Slides column maps each session to the matching slides in the facilitator deck "
          "(facilitator-slides.pptx, 94 slides) so trainers can pace delivery against the deck.")
for nm, theme, rows in [
    ("Day 1 — Foundations & Power Automate", "Workflow concepts + your first five flows", DAY1),
    ("Day 2 — Building Business Agents with Copilot Studio", "Agents, knowledge/RAG, tools, and agent + flow", DAY2),
    ("Day 3 — End-to-End Workflow Automation & Workshop", "Combine agents + flows, then build your own", DAY3),
]:
    heading(doc, nm, 2)
    para(doc, theme + ".")
    schedule_table(doc, rows)
    doc.add_paragraph()

heading(doc, "Tools & Resources", 1)
info_table(doc, [
    ("Microsoft Power Automate", "make.powerautomate.com — build and run flows"),
    ("Microsoft Copilot Studio", "copilotstudio.microsoft.com — build and publish agents"),
    ("Power Platform admin center", "admin.powerplatform.microsoft.com — create the Sandbox environment (Dataverse)"),
    ("Microsoft 365", "Outlook, Excel (OneDrive), Microsoft Forms"),
    ("Learner Guide", "Full step-by-step guide for every lab (provided)"),
])

heading(doc, "Assessment", 1)
para(doc, "Assessment is continuous and practical. Each lab includes a Checkpoint that the facilitator verifies. "
          "The Day-3 capstone (Lab 16) is the summative task: each participant designs, builds, tests, and presents "
          "an original end-to-end workflow against the published quality bar (clear trigger, data logged, at least one "
          "condition, an approval or AI-generated response, notifications for every outcome, and tested happy path + branches).")

add_footer(doc)
prodoc.enable_update_fields(doc)
out = os.path.join(REPO, f"courseware/LP-{TITLE}.docx")
doc.save(out)
print("Day totals:", {n: total(r) for n,r in [("D1",DAY1),("D2",DAY2),("D3",DAY3)]})
print("Wrote", out)
